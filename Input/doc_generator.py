# Document Generated for the updated text
from docx import Document
from docx.shared import Inches

# Final text paragraph
final_text = 'With reasonable cause, client both parties may terminate this Agreement, \
effective immediately upon giving written notice. \n \
Reasonable cause includes: \n \
•   Violation of this Agreement, or \n \
•   Any act exposing liability to others for personal injury or property damage.'

document = Document()

document.add_heading('Final Termination Clause with Feedback', 0)

document.add_heading('Terminating the Agreement', level=1)

p = document.add_paragraph(final_text)

document.save('generated_document.docx')

print('Document Generated!')
